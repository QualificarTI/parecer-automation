"""
main.py — Serviço "gerar_parecer"
-----------------------------------------------------------------
Este serviço NÃO fala mais direto com o SharePoint / Microsoft Graph (a
versão antiga fazia isso e exigia um App Registration no Entra ID com
consentimento de administrador). Agora ele só recebe dados e devolve
arquivos prontos — quem grava no SharePoint é o próprio Power Automate,
usando os conectores nativos de SharePoint (login normal, sem admin).
Fluxo, na ordem:
  1. O flow do Power Automate baixa o Base de Talentos atual do
     SharePoint (ação nativa "Obter conteúdo do arquivo") e converte pra
     base64.
  2. O flow chama este serviço via HTTP POST, mandando no corpo:
       - todos os 138 campos do parecer (o mesmo dados_parecer_json já
         parseado pelo Parse JSON do flow);
       - "excel_atual_base64": o conteúdo do Excel atual, em base64
         (pode vir vazio/omitido se o arquivo ainda não existir no
         SharePoint — nesse caso um Excel novo é criado do zero).
  3. Este serviço gera o Word (docxtpl) e faz o upsert no Excel (mesma
     regra de dedupe do registrar_talentos.py original: chave composta
     Nome_Completo+Cliente+Cargo_Aplicado+Recrutador_Responsavel,
     mantendo a linha com a Data_Entrevista mais recente).
  4. Devolve ao flow, em base64: o Word novo e o Excel atualizado, mais o
     nome do arquivo do parecer e o status da Base de Talentos
     ("appended" / "updated" / "skipped" / "sem_dados_minimos").
  5. O flow grava os dois arquivos no SharePoint com as ações nativas
     "Criar arquivo" / "Atualizar arquivo" e responde ao agente.
Rodar localmente pra testar:
    uvicorn main:app --host 0.0.0.0 --port 8000
Variáveis de ambiente necessárias — ver .env.example
"""
import base64
import io
import os
import re
import unicodedata
import datetime
from pathlib import Path

from fastapi import FastAPI, Header, HTTPException
from docxtpl import DocxTemplate
from docx import Document
from openpyxl import Workbook, load_workbook

# ============================================================
# Config (variáveis de ambiente — ver .env.example)
# ============================================================
# Chave simples que o flow do Power Automate envia no header pra autenticar a chamada
API_KEY = os.environ.get("PARECER_API_KEY", "")

# Nome do arquivo da Base de Talentos devolvido no JSON de resposta (só
# informativo pro flow usar ao gravar — o caminho real dentro do
# SharePoint é definido no próprio flow, não aqui).
EXCEL_FILENAME = os.environ.get("PARECER_EXCEL_FILENAME", "base_talentos.xlsx")

# Pasta local só pra depuração manual (ex.: testar com curl sem o flow) —
# não tem nenhum papel em produção, é só um rastro pra facilitar debug.
LOCAL_OUTPUT_DIR = Path(__file__).parent / "output_local"
TEMPLATE_PATH = Path(__file__).parent / "ModeloParecer_jinja.docx"

BASE_TALENTOS_FIELDS = [
    "Nome_Completo", "Telefone", "Cidade", "UF", "LinkedIn_URL",
    "Cliente", "Cargo_Aplicado", "Senioridade_Aplicada", "Data_Entrevista",
    "Anos_Exp_No_Cargo", "Anos_Exp_Geral", "Recomendacao",
    "Escolaridade", "Area_Formacao", "Recrutador_Responsavel",
    "Palavras_Chave_5", "Observacoes",
]

# Campos do parecer (mesmo FIELD_ORDER do gerar_parecer.py original)
FIELD_ORDER = [
    "Cargo", "Cliente", "Recrutador", "Data_Envio", "Recomendacao",
    "Nome_Candidato", "Tempo_Experiencia",
]
for i in range(1, 16):
    FIELD_ORDER += [f"REQO{i:02d}", f"REQO{i:02d}_Analise", f"REQO{i:02d}_Comprovante"]
for i in range(1, 11):
    FIELD_ORDER += [f"REQD{i:02d}", f"REQD{i:02d}_Analise", f"REQD{i:02d}_Comprovante"]
for i in range(1, 11):
    FIELD_ORDER += [f"CERO{i:02d}", f"CERO{i:02d}_Analise", f"CERO{i:02d}_Comprovante"]
FIELD_ORDER += [
    "Comunicacao", "Comunicacao_Analise", "Comunicacao_Comprovante",
    "Organizacao", "Organizacao_Analise", "Organizacao_Comprovante",
    "Analise_Critica", "Analise_Critica_Analise", "Analise_Critica_Comprovante",
    "Criatividade", "Criatividade_Analise", "Criatividade_Comprovante",
    "Inovacao", "Inovacao_Analise", "Inovacao_Comprovante",
    "Parecer_Final",
]

# Rótulos estáticos das competências — o agente não precisa enviar isso,
# o servidor já preenche sozinho (evita 5 campos redundantes no JSON).
STATIC_LABELS = {
    "Comunicacao": "Comunicação",
    "Organizacao": "Organização",
    "Analise_Critica": "Análise Crítica",
    "Criatividade": "Criatividade",
    "Inovacao": "Inovação",
}

ALL_TEMPLATE_FIELDS = FIELD_ORDER  # os 128 campos que o docxtpl espera

# ============================================================
# Normalização do payload — o agente já manda JSON estruturado (via
# Parse JSON no flow), então aqui só garantimos que as 138 chaves
# existem (preenche com "" o que faltar) e limpamos tipos estranhos.
# ============================================================
def _norm(s: str) -> str:
    s = (s or "").strip().replace("-", "_").replace(" ", "_")
    s = "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")
    return s.lower()


def sanitize_filename(s: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "_", (s or "").strip())


# Caracteres de controle que o padrão XML 1.0 NÃO aceita dentro de texto
# (o Word usa XML por baixo dos panos). Se um valor vindo do agente —
# currículo colado, transcrição de entrevista, etc. — tiver algum desses
# caracteres "invisíveis" (comuns em textos copiados de PDF/Teams/Word),
# o .docx final continua sendo um .zip válido (por isso o "unzip -t"
# não acusava nada), mas o XML interno fica inválido e o Word recusa
# abrir com "Ocorreu um erro no Word ao tentar abrir o ficheiro".
# Faixa válida no XML 1.0: #x9 | #xA | #xD | [#x20-#xD7FF] |
# [#xE000-#xFFFD] | [#x10000-#x10FFFF]. Removemos tudo fora disso.
_XML_INVALID_CHARS_RE = re.compile(
    "[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)


def strip_xml_invalid_chars(s: str) -> str:
    if not s:
        return s
    return _XML_INVALID_CHARS_RE.sub("", s)


def parse_date_flexible(s: str):
    if not s:
        return None
    s = s.strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except Exception:
            pass
    return None


def normalize_payload(raw: dict) -> dict:
    """Garante que todas as 138 chaves esperadas existem, mesmo que o
    agente tenha mandado alguma vazia/faltando. Valores None viram ''.
    Também remove caracteres de controle inválidos em XML (ver
    strip_xml_invalid_chars) — é isso que evita o Word corromper o
    arquivo quando algum campo (currículo/transcrição colados) traz
    caracteres invisíveis que o XML do .docx não aceita."""
    all_keys = FIELD_ORDER + [k for k in BASE_TALENTOS_FIELDS if k not in FIELD_ORDER]
    return {
        k: strip_xml_invalid_chars("" if raw.get(k) is None else str(raw.get(k)))
        for k in all_keys
    }


# ============================================================
# Excel — mesma lógica de upsert do registrar_talentos.py original
# ============================================================
# Larguras de coluna (em "caracteres", unidade padrão do openpyxl/Excel).
# Aplicadas sempre que a planilha passa por aqui, então mesmo uma
# base_talentos.xlsx antiga (já criada com colunas estreitas) volta a
# ficar com larguras razoáveis a cada parecer novo.
_COLUMN_WIDTHS = {
    "Nome_Completo": 28,
    "Telefone": 16,
    "Cidade": 18,
    "UF": 6,
    "LinkedIn_URL": 32,
    "Cliente": 18,
    "Cargo_Aplicado": 26,
    "Senioridade_Aplicada": 18,
    "Data_Entrevista": 14,
    "Anos_Exp_No_Cargo": 12,
    "Anos_Exp_Geral": 12,
    "Recomendacao": 16,
    "Escolaridade": 20,
    "Area_Formacao": 22,
    "Recrutador_Responsavel": 20,
    "Palavras_Chave_5": 30,
    "Observacoes": 60,
}


def ensure_excel_headers(ws):
    if ws.max_row == 1 and all(ws.cell(1, c + 1).value is None for c in range(len(BASE_TALENTOS_FIELDS))):
        for i, h in enumerate(BASE_TALENTOS_FIELDS, start=1):
            ws.cell(1, i, h)
    for i, h in enumerate(BASE_TALENTOS_FIELDS, start=1):
        width = _COLUMN_WIDTHS.get(h)
        if width:
            ws.column_dimensions[ws.cell(1, i).column_letter].width = width


def _columns_map():
    return {h: i for i, h in enumerate(BASE_TALENTOS_FIELDS, start=1)}


def _row_key_tuple(row_dict: dict):
    return (
        _norm(row_dict.get("Nome_Completo", "")),
        _norm(row_dict.get("Cliente", "")),
        _norm(row_dict.get("Cargo_Aplicado", "")),
        _norm(row_dict.get("Recrutador_Responsavel", "")),
    )


def _find_matching_row(ws, row_dict: dict):
    key = _row_key_tuple(row_dict)
    if any(not part for part in key):
        return None
    cols = _columns_map()
    max_row = ws.max_row or 1
    for r in range(2, max_row + 1):
        k2 = (
            _norm(ws.cell(r, cols["Nome_Completo"]).value or ""),
            _norm(ws.cell(r, cols["Cliente"]).value or ""),
            _norm(ws.cell(r, cols["Cargo_Aplicado"]).value or ""),
            _norm(ws.cell(r, cols["Recrutador_Responsavel"]).value or ""),
        )
        if k2 == key:
            return r
    return None


def _is_newer(new_date_str: str, old_date_str: str) -> bool:
    nd = parse_date_flexible(new_date_str)
    od = parse_date_flexible(old_date_str)
    if nd and od:
        return nd >= od
    return True


def upsert_talentos(existing_bytes: bytes | None, row: dict) -> tuple[bytes, str]:
    if existing_bytes:
        wb = load_workbook(io.BytesIO(existing_bytes))
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
    ensure_excel_headers(ws)
    cols = _columns_map()
    r_match = _find_matching_row(ws, row)
    if r_match is None:
        ws.append([row.get(k, "") for k in BASE_TALENTOS_FIELDS])
        status = "appended"
    else:
        old_date = ws.cell(r_match, cols["Data_Entrevista"]).value or ""
        new_date = row.get("Data_Entrevista", "")
        if _is_newer(new_date, old_date):
            for i, field in enumerate(BASE_TALENTOS_FIELDS, start=1):
                ws.cell(r_match, i, row.get(field, ""))
            status = "updated"
        else:
            status = "skipped"
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue(), status


# ============================================================
# Geração do Word via docxtpl
# ============================================================
# Índices das tabelas do template que têm linhas "opcionais" (slots de
# REQO/REQD/CERO que nem sempre são todos usados pela vaga). As duas
# primeiras linhas de cada uma são cabeçalho (título do grupo + nomes das
# colunas) — os dados começam na linha de índice 2. Depende da estrutura
# atual do ModeloParecer_jinja.docx (tabela 2 = Requisitos Obrigatórios,
# 3 = Desejáveis, 4 = Certificações) — se o template mudar de estrutura,
# esses índices precisam ser revisados.
REQUIREMENT_TABLE_INDEXES = [2, 3, 4]
HEADER_ROWS_PER_TABLE = 2

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _row_text(tr) -> str:
    """Extrai todo o texto de uma linha de tabela (<w:tr>), incluindo texto
    dentro de células envolvidas em controles de conteúdo (<w:sdt>). O
    python-docx padrão (row.cells) não enxerga células assim — só olha
    filhos diretos de <w:tr> — por isso precisamos ler o XML na mão."""
    return "".join(t.text or "" for t in tr.iter(_W + "t"))


def remove_empty_requirement_rows(doc):
    """Remove linhas de tabela cujos campos ficaram todos vazios (slots de
    REQO/REQD/CERO não usados pela vaga) — evita linhas em branco feias no
    documento final."""
    for t_idx in REQUIREMENT_TABLE_INDEXES:
        if t_idx >= len(doc.tables):
            continue
        table = doc.tables[t_idx]
        all_trs = table._tbl.findall(_W + "tr")
        data_trs = all_trs[HEADER_ROWS_PER_TABLE:]
        for tr in data_trs:
            if _row_text(tr).strip() == "":
                tr.getparent().remove(tr)


def remove_empty_requirement_blocks(doc):
    """Remove a tabela inteira (título + cabeçalho de colunas, ambos em
    azul) quando nenhuma linha de dado sobrou nela — ex.: vaga sem nenhum
    requisito desejável ou sem nenhuma certificação obrigatória. Precisa
    rodar DEPOIS de remove_empty_requirement_rows, senão nunca vai achar
    tabela "vazia" (as linhas em branco ainda estariam lá)."""
    target_tables = [doc.tables[i] for i in REQUIREMENT_TABLE_INDEXES if i < len(doc.tables)]
    for table in target_tables:
        remaining_trs = table._tbl.findall(_W + "tr")
        if len(remaining_trs) <= HEADER_ROWS_PER_TABLE:
            table._tbl.getparent().remove(table._tbl)


def render_docx(payload: dict) -> bytes:
    ctx = {f: payload.get(f, "") for f in ALL_TEMPLATE_FIELDS}
    ctx.update(STATIC_LABELS)  # garante os rótulos das competências
    tpl = DocxTemplate(str(TEMPLATE_PATH))
    tpl.render(ctx)
    buf = io.BytesIO()
    tpl.save(buf)
    data = buf.getvalue()

    # Autoverificação: tenta reabrir o .docx recém-gerado antes de devolvê-lo.
    # Um .zip pode estar tecnicamente íntegro (unzip -t não acusa nada) e
    # mesmo assim ter XML interno inválido, que o Word recusa abrir. Se
    # isso acontecer aqui, preferimos falhar com um erro 500 claro do que
    # devolver pro flow um arquivo que só vai quebrar na hora de abrir no
    # Word.
    try:
        doc = Document(io.BytesIO(data))
        _ = len(doc.paragraphs)  # força o parse completo do document.xml
    except Exception as e:
        raise RuntimeError(
            f"O .docx gerado não passou na verificação de integridade "
            f"(provavelmente algum campo do payload contém caracteres "
            f"inválidos para XML): {e}"
        )

    # Remove linhas vazias (slots de requisito/certificação não usados) e,
    # se a tabela inteira ficou sem nenhuma linha de dado, remove o bloco
    # inteiro (título + cabeçalho). Depois resalva e reverifica de novo
    # pra garantir que a limpeza não quebrou nada.
    remove_empty_requirement_rows(doc)
    remove_empty_requirement_blocks(doc)
    buf2 = io.BytesIO()
    doc.save(buf2)
    data = buf2.getvalue()
    try:
        Document(io.BytesIO(data))
    except Exception as e:
        raise RuntimeError(f"Falha ao limpar linhas vazias do documento: {e}")

    return data


# ============================================================
# API
# ============================================================
app = FastAPI(title="Gerar Parecer", version="2.0")


@app.post("/gerar-parecer")
def gerar_parecer(body: dict, x_api_key: str = Header(default="")):
    """Recebe o JSON de 138 campos do parecer + (opcional)
    excel_atual_base64 com o conteúdo atual da Base de Talentos. Devolve
    o Word gerado e o Excel atualizado, ambos em base64, prontos pro
    flow gravar no SharePoint com as ações nativas."""
    if API_KEY and x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="API key inválida.")

    if not body:
        raise HTTPException(status_code=400, detail="Corpo da requisição vazio.")

    payload = normalize_payload(body)

    # nome do arquivo — usa a data de REALIZAÇÃO da entrevista
    # (Data_Entrevista), não a data de envio, pra bater com o padrão já
    # usado hoje (ex.: "Parecer - ANS - Analista de Testes - Fulano -
    # 2026-05-07.docx"). Cai pra Data_Envio ou hoje se não vier.
    hoje = datetime.date.today()
    data_realizacao = (
        parse_date_flexible(payload.get("Data_Entrevista"))
        or parse_date_flexible(payload.get("Data_Envio"))
        or hoje
    )
    cliente = sanitize_filename(payload.get("Cliente"))
    cargo = sanitize_filename(payload.get("Cargo"))
    cand = sanitize_filename(payload.get("Nome_Candidato"))
    dt_str = data_realizacao.strftime("%Y-%m-%d")
    filename = f"Parecer - {cliente} - {cargo} - {cand} - {dt_str}.docx"

    # 1) Gera o docx
    try:
        docx_bytes = render_docx(payload)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Falha ao gerar o documento Word: {e}")

    # 2) Upsert no Excel (recebido em base64, se houver)
    tem_dados_base = any(
        (payload.get(k) or "").strip() for k in ("Nome_Completo", "LinkedIn_URL", "Telefone")
    )
    excel_base64_out = ""
    base_status = "sem_dados_minimos"
    if tem_dados_base:
        excel_atual_b64 = (body.get("excel_atual_base64") or "").strip()
        existing_bytes = None
        if excel_atual_b64:
            try:
                existing_bytes = base64.b64decode(excel_atual_b64)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"excel_atual_base64 inválido: {e}")
        try:
            new_excel_bytes, base_status = upsert_talentos(existing_bytes, payload)
            excel_base64_out = base64.b64encode(new_excel_bytes).decode("ascii")
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Falha ao atualizar a Base de Talentos: {e}")

    # 3) Cópia local só pra depuração manual (não afeta a resposta ao flow)
    try:
        LOCAL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        (LOCAL_OUTPUT_DIR / filename).write_bytes(docx_bytes)
        if excel_base64_out:
            (LOCAL_OUTPUT_DIR / EXCEL_FILENAME).write_bytes(base64.b64decode(excel_base64_out))
    except Exception:
        pass

    return {
        "status": "sucesso",
        "documento_nome": filename,
        "documento_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "excel_nome": EXCEL_FILENAME,
        "excel_base64": excel_base64_out,
        "base_talentos_status": base_status,
    }


@app.get("/health")
def health():
    return {"status": "ok", "modo": "arquivos via base64 (sem Graph/SharePoint direto)"}
